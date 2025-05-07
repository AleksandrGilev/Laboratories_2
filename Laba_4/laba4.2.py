import csv
from docxtpl import DocxTemplate
from docx import Document


def create_template():
    """Создаем шаблон документа для отчета о марафонах"""
    doc = Document()

    p = doc.add_paragraph()
    p.add_run("{% for year_data in years_data %}")

    h1 = doc.add_paragraph()
    h1.add_run("Год проведения марафона: {{ year_data.year }}").bold = True

    doc.add_paragraph("{% for marathon in year_data.marathons %}")

    city = doc.add_paragraph()
    city.add_run("Город, где проводился марафон: {{ marathon.city }}").bold = True

    doc.add_paragraph(
        "Имя победителя-мужчины: {{ marathon.male_winner }}, время на дистанции: {{ marathon.male_time }}")
    doc.add_paragraph(
        "Имя победителя-женщины: {{ marathon.female_winner }}, время на дистанции: {{ marathon.female_time }}")

    doc.add_paragraph("-" * 40)

    doc.add_paragraph("{% endfor %}")

    doc.add_paragraph("{% if not loop.last %}")
    doc.add_page_break()
    doc.add_paragraph("{% endif %}")

    doc.add_paragraph("{% endfor %}")

    doc.save("template.docx")

    print("Шаблон template.docx создан")
    return "template.docx"


def main():
    template_path = create_template()

    doc = DocxTemplate(template_path)

    with open("data_marathon.csv", "r", encoding="utf-8") as file:
        row_data = list(csv.DictReader(file))

    print("Загружено записей:", len(row_data))

    # Словарь для группировки данных по годам
    data_by_year = {}

    for item in row_data:
        year = item["year"]
        city = item["city"]
        gender = item["gender"]

        if year not in data_by_year:
            data_by_year[year] = {}

        if city not in data_by_year[year]:
            data_by_year[year][city] = {"Male": None, "Female": None}

        # Проверяем лучшее время и обновляем запись
        if data_by_year[year][city][gender] is None or item["time"] < data_by_year[year][city][gender]["time"]:
            data_by_year[year][city][gender] = {
                "name": item["name"],
                "time": item["time"],
                "country": item["country"]
            }

    years_data = []

    for year in sorted(data_by_year.keys()):
        marathon_cities = []

        # Добавляем данные о каждом городе в этом году
        for city in sorted(data_by_year[year].keys()):
            city_data = {
                "city": city,
                "male_winner": data_by_year[year][city]["Male"]["name"] if data_by_year[year][city]["Male"] else "",
                "male_time": data_by_year[year][city]["Male"]["time"] if data_by_year[year][city]["Male"] else "",
                "female_winner": data_by_year[year][city]["Female"]["name"] if data_by_year[year][city][
                    "Female"] else "",
                "female_time": data_by_year[year][city]["Female"]["time"] if data_by_year[year][city]["Female"] else ""
            }
            marathon_cities.append(city_data)

        year_data = {
            "year": year,
            "marathons": marathon_cities
        }
        years_data.append(year_data)

    context = {
        "years_data": years_data
    }

    # Рендерим и сохраняем
    doc.render(context)
    doc.save("result.docx")
    print(
        f"Отчет сохранен в файл result.docx. Обработано {len(years_data)} лет и {sum(len(y['marathons']) for y in years_data)} марафонов.")


if __name__ == "__main__":
    main()
