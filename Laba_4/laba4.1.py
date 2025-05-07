import PIL.ImageQt
from docx import Document
from docx.shared import Mm, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = Document()

doc.add_heading('Автобиография')
doc.add_heading('Школа', level=2)

p1 = doc.add_paragraph(
    'Родился в Архангельской области, городе Котлас, 15 мая 2006 года.'
    'Первые 7 лет жил в пгт Вычегодский. В конце 1 класса переехал на ДОК.'
    'Первые 9 классов учился в МОУ СОШ №18, 10 и 11 классы учился в МОУ ОЛ №3.'
)

p1.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

p2 = doc.add_paragraph(
    'Закончил музыкальную школу, занимался легкой атлетикой и MMA.'
    'Умею играть на гитаре и фортепиано, в городе ходил с концертами в школах, колледжах, Русско-немецкий центр, Котласский КЦСО.'
)

p2.paragraph_format.first_line_indent = Mm(30)
p2.paragraph_format.left_indent = Mm(40)
p2.paragraph_format.right_indent = Mm(-50)
p2.paragraph_format.space_before = Mm(30)
p2.paragraph_format.line_spacing = Mm(10)

picture = doc.add_paragraph()
picture_run = picture.add_run()
picture_run.add_picture('для лаб.jpg', width=Mm(100), height=Mm(100))
picture.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.save('res.docx')
